import re
import os
import tkinter as tk
from tkinter import filedialog
from docx import Document

def clean_all_stl_files_in_folder(folder_path, output_filename):
    # Define the regex patterns
    timestamp_pattern = r'^\d{2}:\d{2}:\d{2}:\d{2},\d{2}:\d{2}:\d{2}:\d{2},\s*'  # Matches timestamps
    comment_pattern = r'^//.*'  # Matches lines that start with //

    # Create a new Word Document
    document = Document()

    # Loop through all .stl files in the specified folder
    for idx, filename in enumerate(os.listdir(folder_path)):
        if filename.endswith('.stl'):
            input_filepath = os.path.join(folder_path, filename)

            # Add the name of the input file at the beginning of each section in the document
            document.add_paragraph(f"Subtitle File: {filename}\n")

            # Read the STL subtitle file
            with open(input_filepath, 'r', encoding='utf-8') as file:
                lines = file.readlines()

            # Process each line of the STL file
            for line in lines:
                # Skip lines that are comments (start with "//")
                if re.match(comment_pattern, line):
                    continue

                # Remove timestamp if present and clean up formatting
                cleaned_line = re.sub(timestamp_pattern, '', line)
                cleaned_line = cleaned_line.replace('<br>', ' ').strip()
                if cleaned_line:  # Only add non-empty lines
                    document.add_paragraph(cleaned_line)

            # Insert a page break after each file, except the last one
            if idx < len(os.listdir(folder_path)) - 1:
                document.add_page_break()

    # Save the cleaned-up text into a single Word file
    document.save(output_filename)
    print(f"All cleaned subtitles saved to '{output_filename}'")

def browse_folder():
    folder_path = filedialog.askdirectory()  # Opens a dialog to choose a folder
    if folder_path:
        output_word_file = os.path.join(folder_path, 'compiled_cleaned_subs.docx')
        clean_all_stl_files_in_folder(folder_path, output_word_file)

# Create a simple GUI
root = tk.Tk()
root.title("STL Subtitle Cleaner")

# Set up the window size
root.geometry("400x200")

# Add a label to guide the user
label = tk.Label(root, text="Select a folder containing .stl files to process", pady=20)
label.pack()

# Add a button that will allow the user to select the folder
button = tk.Button(root, text="Browse Folder", command=browse_folder, padx=10, pady=5)
button.pack()

# Run the tkinter main loop
root.mainloop()
